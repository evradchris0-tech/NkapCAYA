"""
Génère le devis CAYA au format Word (.docx).

Pré-requis :
    pip install python-docx

Exécution :
    python generate_devis.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Données du devis ────────────────────────────────────────────────

EMETTEUR = {
    "nom": "OMGBA FOUDA Christ",
    "activite": "Freelance — Développement logiciel",
    "email": "chrisomgba04@gmail.com",
    "telephone": "+237 655 74 67 14",
    "ville": "Yaoundé, Cameroun",
}

CLIENT = {
    "nom": "Club des Amis de Yaoundé (CAYA)",
    "type": "Association de tontine",
    "ville": "Yaoundé, Cameroun",
}

DEVIS = {
    "numero": "DEV-CAYA-2026-001",
    "date_emission": "01 mai 2026",
    "periode": "Décembre 2025 — Mai 2026 (≈ 5 mois)",
    "validite": "30 jours",
}

DESCRIPTION = (
    "Conception et développement d'une plateforme web complète de gestion de tontine "
    "pour le CAYA — application greenfield couvrant 10 modules métiers (membres, "
    "exercices fiscaux, sessions mensuelles, épargne, prêts, caisse de secours, "
    "bénéficiaires, cassation, rapports), avec pipeline de déploiement automatisé."
)

PRESTATIONS = [
    ("Analyse & conception (UML, modèle de données, architecture en couches)", 800_000, 5_000),
    ("Développement Backend NestJS — 10 modules métier + Prisma ORM", 2_500_000, 18_000),
    ("Développement Frontend Next.js — dashboards, formulaires, navigation", 2_000_000, 15_000),
    ("Module Rapports & Exports CAYABASE (Excel multi-feuilles, PDF)", 600_000, 5_000),
    ("Module Import d'exercice depuis fichier Excel CAYABASE", 400_000, 3_000),
    ("Sécurité — Authentification JWT, RBAC, journalisation des actions", 500_000, 4_000),
    ("Pipeline CI/CD & déploiement Hostinger automatisé", 400_000, 3_000),
    ("Tests, recette utilisateur & corrections", 600_000, 5_000),
    ("Documentation technique & guide utilisateur", 300_000, 3_000),
    ("Formation administrateurs (2 sessions)", 200_000, 4_000),
    ("Maintenance & évolutions sur 5 mois", 700_000, 5_000),
]

CONDITIONS_PAIEMENT = [
    "Règlement par Mobile Money ou virement bancaire"
]

# ── Couleurs ────────────────────────────────────────────────────────

COLOR_PRIMARY = RGBColor(0x1E, 0x40, 0xAF)   # Bleu CAYA
COLOR_DARK = RGBColor(0x1F, 0x29, 0x37)
COLOR_GRAY = RGBColor(0x6B, 0x72, 0x80)
COLOR_GREEN = RGBColor(0x05, 0x96, 0x69)
COLOR_HEADER_BG = "1E40AF"
COLOR_TOTAL_BG = "F3F4F6"
COLOR_DISCOUNT_BG = "ECFDF5"


# ── Helpers ─────────────────────────────────────────────────────────

def fmt_fcfa(n: int) -> str:
    return f"{n:,}".replace(",", " ") + " FCFA"


def set_cell_bg(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for border in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{border}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "D1D5DB")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_heading(doc, text, size=14, color=COLOR_PRIMARY, bold=True, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    return p


def add_paragraph(doc, text, size=10, color=COLOR_DARK, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return p


# ── Génération du document ──────────────────────────────────────────

def generate_devis(output_path: str = "devis_caya.docx"):
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Police par défaut
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ── Titre ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DEVIS")
    run.font.size = Pt(28)
    run.font.color.rgb = COLOR_PRIMARY
    run.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"N° {DEVIS['numero']}")
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_GRAY

    doc.add_paragraph()

    # ── En-tête (Émetteur / Client) ──
    info_table = doc.add_table(rows=1, cols=2)
    info_table.autofit = False
    info_table.columns[0].width = Cm(8)
    info_table.columns[1].width = Cm(8)

    cell_em = info_table.rows[0].cells[0]
    p = cell_em.paragraphs[0]
    run = p.add_run("ÉMETTEUR\n")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_PRIMARY
    cell_em.add_paragraph(EMETTEUR["nom"]).runs[0].bold = True
    cell_em.add_paragraph(EMETTEUR["activite"])
    cell_em.add_paragraph(f"📧 {EMETTEUR['email']}")
    cell_em.add_paragraph(f"📞 {EMETTEUR['telephone']}")
    cell_em.add_paragraph(EMETTEUR["ville"])

    cell_cl = info_table.rows[0].cells[1]
    p = cell_cl.paragraphs[0]
    run = p.add_run("CLIENT\n")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_PRIMARY
    cell_cl.add_paragraph(CLIENT["nom"]).runs[0].bold = True
    cell_cl.add_paragraph(CLIENT["type"])
    cell_cl.add_paragraph(CLIENT["ville"])

    doc.add_paragraph()

    # ── Bloc infos devis ──
    info_devis = doc.add_table(rows=4, cols=2)
    info_devis.autofit = False
    info_devis.columns[0].width = Cm(5)
    info_devis.columns[1].width = Cm(11)

    rows_data = [
        ("Date d'émission", DEVIS["date_emission"]),
        ("Période concernée", DEVIS["periode"]),
        ("Validité du devis", DEVIS["validite"]),
        ("Devise", "Franc CFA (XAF)"),
    ]
    for i, (label, value) in enumerate(rows_data):
        c1 = info_devis.rows[i].cells[0]
        c2 = info_devis.rows[i].cells[1]
        c1.paragraphs[0].add_run(label).bold = True
        c2.paragraphs[0].add_run(value)
        set_cell_bg(c1, "F9FAFB")
        set_cell_borders(c1)
        set_cell_borders(c2)

    doc.add_paragraph()

    # ── Description ──
    add_heading(doc, "Description du projet", size=12)
    add_paragraph(doc, DESCRIPTION, size=10)

    # ── Tableau prestations ──
    add_heading(doc, "Détail des prestations", size=12, space_before=18)

    presta_table = doc.add_table(rows=1, cols=4)
    presta_table.autofit = False
    presta_table.columns[0].width = Cm(0.8)
    presta_table.columns[1].width = Cm(8.5)
    presta_table.columns[2].width = Cm(3.5)
    presta_table.columns[3].width = Cm(3.5)

    # En-tête
    headers = ["#", "Prestation", "Tarif entreprise", "Tarif consenti CAYA"]
    hdr_cells = presta_table.rows[0].cells
    for i, header in enumerate(headers):
        cell = hdr_cells[i]
        set_cell_bg(cell, COLOR_HEADER_BG)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        if i in (2, 3):
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif i == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Lignes
    total_entreprise = 0
    total_consenti = 0
    for idx, (label, prix_ent, prix_consenti) in enumerate(PRESTATIONS, start=1):
        row = presta_table.add_row().cells
        for c in row:
            set_cell_borders(c)

        row[0].paragraphs[0].add_run(str(idx)).bold = True
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        row[1].paragraphs[0].add_run(label).font.size = Pt(9)

        p2 = row[2].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p2.add_run(fmt_fcfa(prix_ent))
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_GRAY

        p3 = row[3].paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p3.add_run(fmt_fcfa(prix_consenti))
        run.font.size = Pt(9)
        run.bold = True
        run.font.color.rgb = COLOR_GREEN

        total_entreprise += prix_ent
        total_consenti += prix_consenti

    # Ligne TOTAL HT
    row = presta_table.add_row().cells
    for c in row:
        set_cell_borders(c)
        set_cell_bg(c, COLOR_TOTAL_BG)
    row[0].merge(row[1])
    p = row[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("TOTAL HT")
    run.bold = True
    run.font.size = Pt(10)

    p = row[2].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(fmt_fcfa(total_entreprise))
    run.bold = True
    run.font.size = Pt(10)

    p = row[3].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(fmt_fcfa(total_consenti))
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_GREEN

    # Ligne TVA
    row = presta_table.add_row().cells
    for c in row:
        set_cell_borders(c)
    row[0].merge(row[1])
    p = row[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("TVA (non applicable — auto-entrepreneur)")
    run.font.size = Pt(9)
    run.italic = True
    row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    row[2].paragraphs[0].add_run("0 FCFA").font.size = Pt(9)
    row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    row[3].paragraphs[0].add_run("0 FCFA").font.size = Pt(9)

    # Ligne TOTAL TTC
    row = presta_table.add_row().cells
    for c in row:
        set_cell_borders(c)
        set_cell_bg(c, COLOR_HEADER_BG)
    row[0].merge(row[1])
    p = row[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("TOTAL TTC")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p = row[2].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(fmt_fcfa(total_entreprise))
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p = row[3].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(fmt_fcfa(total_consenti))
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # ── Bloc remise ──
    doc.add_paragraph()
    remise = total_entreprise - total_consenti
    pct = (remise / total_entreprise) * 100
    remise_table = doc.add_table(rows=1, cols=1)
    cell = remise_table.rows[0].cells[0]
    set_cell_bg(cell, COLOR_DISCOUNT_BG)
    set_cell_borders(cell)
    p = cell.paragraphs[0]
    run = p.add_run(f"💚 Remise consentie au CAYA : {fmt_fcfa(remise)}  ({pct:.1f} %)")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_GREEN
    p2 = cell.add_paragraph()
    run2 = p2.add_run(
        "Geste de soutien à l'association — le tarif consenti vise uniquement "
        "à amortir une partie du temps investi depuis le démarrage du projet."
    )
    run2.font.size = Pt(9)
    run2.italic = True
    run2.font.color.rgb = COLOR_DARK

    # ── Conditions de paiement ──
    add_heading(doc, "Conditions de paiement", size=12, space_before=18)
    for cond in CONDITIONS_PAIEMENT:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(cond).font.size = Pt(10)

    # ── Signature ──
    doc.add_paragraph()
    doc.add_paragraph()
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.autofit = False
    sig_table.columns[0].width = Cm(8)
    sig_table.columns[1].width = Cm(8)

    c_em = sig_table.rows[0].cells[0]
    c_em.paragraphs[0].add_run("Le prestataire").bold = True
    c_em.add_paragraph(EMETTEUR["nom"])
    c_em.add_paragraph("\n\n_____________________")
    c_em.add_paragraph("Date & signature").runs[0].italic = True

    c_cl = sig_table.rows[0].cells[1]
    c_cl.paragraphs[0].add_run("Le client (Bon pour accord)").bold = True
    c_cl.add_paragraph(CLIENT["nom"])
    c_cl.add_paragraph("\n\n_____________________")
    c_cl.add_paragraph("Date, signature & cachet").runs[0].italic = True

    # ── Footer ──
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        f"Devis émis le {DEVIS['date_emission']} — Valable {DEVIS['validite']}"
    )
    run.font.size = Pt(8)
    run.italic = True
    run.font.color.rgb = COLOR_GRAY

    # Sauvegarde
    doc.save(output_path)
    print(f"Devis genere : {output_path}")
    print(f"  Total entreprise  : {fmt_fcfa(total_entreprise)}")
    print(f"  Total consenti    : {fmt_fcfa(total_consenti)}")
    print(f"  Remise            : {fmt_fcfa(remise)} ({pct:.1f} %)")


if __name__ == "__main__":
    generate_devis()
